#!/bin/python3

from docx import Document
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn




ALL_WORDS1 = [
"apartment",
"apartment block",
"apartment building",
"apartment house",
"attics",
"backyard",
"bar",
"block of flats",
"box-house",
"box junction",
"bungalow",
"club",
"condominium",
"dormitory ",
"duplex",
"entryway",
"estate",
"false-front building",
"flat",
"foyer",
"front room",
"garden",
"grocery store",
"hall",
"hallway",
"honky-tonk",
"housing project",
"living room",
"lobby",
"lodge",
"loft apartment",
"lounge",
"lounge bar",
"lunchroom",
"mobile home",
"parlour",
"porch",
"public house",
"pub",
"ranch house",
"rooming house",
"row house",
"saloon",
"semi-skyscraper",
"shop",
"store",
"subdivision",
"sun lounge",
"sun porch",
"tavern",
"terraced house",
"townhouse",
"yard",
]

ALL_WORDS3 = [
"barman",
"barmaid",
"bartender",
"cab driver",
"cabbie",
"caregiver",
"caretaker",
"carer",
"chaperon",
"hobo",
"janitor",
"porter",
"professor",
"solicitor",
]

ALL_WORDS2 = [
"auto shop",
"avenue",
"beltway",
"blacktop",
"block",
"boulevard",
"bus",
"bus station",
"bypass",
"byroad",
"byway",
"car park",
"caravan",
"come down the pike",
"couch",
"curb",
"dead end",
"detour",
"dirt road",
"dirt track",
"divided highway",
"driveway",
"dual carriageway",
"estate car",
"expressway",
"fender",
"filling station",
"fly-over",
"footpath",
"freeway",
"garage",
"gas station",
"highroad",
"highway",
"hold-up",
"interstate",
"intersection",
"main road",
"main street",
"motorway",
"off-ramp",
"overpass",
"parking lot",
"parkway",
"pavement",
"public highway",
"rat run",
"relief road",
"right up/down someone’s alley",
"right up/down someone’s street",
"ring road",
"road diversion",
"roadhouse",
"rotary",
"roundabout",
"saloon car",
"schedule",
"sedan",
"service road",
"shoulder",
"sidewalk",
"slip road",
"subway",
"superhighway",
"the m25",
"thruway",
"tie-up",
"tarmac",
"tollway",
"traffic circle",
"trail",
"trailer",
"tram",
"trolley",
"trunk road",
"turnpike",
"underground",
"verge",
]


BNC1 = {'apartment': '1853', 'apartment block': '4', 'apartment house': '3', 'attics': '58', 'backyard': '152', 'bar': '10054', 'bungalow': '718', 'club': '19197', 'condominium': '19', 'duplex': '92', 'estate': '7279', 'flat': '10021', 'foyer': '374', 'front room': '7', 'garden': '14388', 'hall': '12410', 'hallway': '418', 'living room': '358', 'lobby': '1117', 'lodge': '1593', 'lounge': '1548', 'lounge bar': '13', 'parlour': '566', 'porch': '515', 'public house': '11', 'pub': '4908', 'saloon': '662', 'shop': '15561', 'store': '6670', 'subdivision': '255', 'sun lounge': '1', 'tavern': '421', 'townhouse': '19', 'yard': '6578'}


COCA1 = {
'apartment': '57674',
'apartment block': '146',
'apartment building': '3035',
'apartment house': '352',
'attics': '389',
'backyard': '11870',
'bar': '75501',
'block of flats': '40',
'bungalow': '1793',
'club': '84284',
'condominium': '1371',
'dormitory ': '1492',
'duplex': '895',
'entryway': '1019',
'estate': '39857',
'flat': '39902',
'foyer': '2709',
'front room': '919',
'garden': '56342',
'grocery store': '7108',
'hall': '78109',
'hallway': '13481',
'honky-tonk': '316',
'housing project': '902',
'living room': '23314',
'lobby': '15962',
'lodge': '8254',
'loft apartment': '44',
'lounge': '7772',
'lounge bar': '14',
'lunchroom': '576',
'mobile home': '949',
'parlour': '366',
'porch': '17236',
'public house': '128',
'pub': '7013',
'ranch house': '634',
'rooming house': '262',
'row house': '267',
'saloon': '2245',
'shop': '52849',
'store': '97612',
'subdivision': '3027',
'sun lounge': '20',
'sun porch': '84',
'tavern': '3823',
'terraced house': '15',
'townhouse': '1254',
'yard': '29772',
}


BNC2 = {'avenue': '1866', 'beltway': '7', 'blacktop': '6', 'block': '6498', 'boulevard': '279', 'bus': '5349', 'bus station': '3', 'bypass': '733', 'byroad': '14', 'byway': '54', 'car park': '120', 'caravan': '1272', 'couch': '481', 'curb': '680', 'dead end': '77', 'detour': '244', 'dirt track': '8', 'driveway': '233', 'dual carriageway': '7', 'estate car': '12', 'expressway': '38', 'fender': '578', 'filling station': '4', 'footpath': '676', 'freeway': '53', 'garage': '2303', 'highway': '1383', 'hold-up': '73', 'interstate': '66', 'intersection': '303', 'main road': '2', 'main street': '4', 'motorway': '1392', 'parking lot': '13', 'parkway': '47', 'pavement': '1718', 'rat run': '3', 'ring road': '14', 'roadhouse': '12', 'rotary': '311', 'roundabout': '675', 'schedule': '3104', 'sedan': '45', 'shoulder': '8392', 'sidewalk': '101', 'slip road': '13', 'subway': '162', 'superhighway': '9', 'tie-up': '49', 'tarmac': '385', 'trail': '1652', 'trailer': '645', 'tram': '765', 'trolley': '853', 'trunk road': '2', 'turnpike': '115', 'underground': '2220', 'verge': '757'}

COCA2 = {
'auto shop': '181',
'avenue': '29262',
'beltway': '2239',
'blacktop': '731',
'block': '62504',
'boulevard': '6856',
'bus': '49900',
'bus station': '1011',
'bypass': '5060',
'byroad': '3',
'byway': '236',
'car park': '303',
'caravan': '2725',
'come down the pike': '51',
'couch': '20472',
'curb': '8893',
'dead end': '2020',
'detour': '1990',
'dirt road': '1908',
'dirt track': '252',
'divided highway': '31',
'driveway': '9841',
'dual carriageway': '13',
'estate car': '3',
'expressway': '1076',
'fender': '1429',
'filling station': '232',
'fly-over': '86',
'footpath': '386',
'freeway': '5494',
'garage': '20334',
'gas station': '4170',
'highroad': '30',
'highway': '29362',
'hold-up': '164',
'interstate': '9371',
'intersection': '7435',
'main road': '1219',
'main street': '6247',
'motorway': '261',
'off-ramp': '201',
'overpass': '825',
'parking lot': '15807',
'parkway': '4486',
'pavement': '5543',
'public highway': '45',
'rat run': '7',
'relief road': '1',
'ring road': '121',
'roadhouse': '444',
'rotary': '2131',
'roundabout': '1012',
'saloon car': '8',
'schedule': '35727',
'sedan': '3816',
'service road': '188',
'shoulder': '48800',
'sidewalk': '11441',
'slip road': '14',
'subway': '9773',
'superhighway': '574',
'the m25': '17',
'thruway': '171',
'tie-up': '92',
'tarmac': '1453',
'tollway': '218',
'traffic circle': '135',
'trail': '37646',
'trailer': '14084',
'tram': '1217',
'trolley': '1693',
'trunk road': '24',
'turnpike': '1129',
'underground': '17229',
'verge': '6797',
}

BNC3 = {'cab driver': '12', 'cabbie': '72', 'caregiver': '12', 'caretaker': '441'}

COCA3 = {
'barman': '394',
'barmaid': '291',
'bartender': '5660',
'cab driver': '915',
'cabbie': '687',
'caregiver': '2914',
'caretaker': '2199',
'carer': '261',
'chaperon': '125',
'hobo': '933',
'janitor': '2336',
'porter': '9144',
'professor': '84811',
'solicitor': '1520',
}


def create_table(document, bnc, coca, all_words):
    #this code creates a table with 2 rows and 2 columens
    table = document.add_table(rows = 1, cols = 4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # table.autofit = False
    # table.allow_autofit = False

    #adding headers rows
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Lexeme'
    hdr_cells[2].text = 'COCA'
    hdr_cells[3].text = 'BNC'
    
    hdr_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
    
    count = 1

    #for word in sorted(set(list(bnc.keys()) + list(coca.keys()))):
    for word in sorted(all_words):
        #second riw
        row_cells = table.add_row().cells
        row_cells[0].text = str(count)
        row_cells[1].text = word
        
        if word in bnc:
            row_cells[3].text = bnc[word]
        
        if word in coca:
            row_cells[2].text = coca[word]
    
        row_cells[0].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[1].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[2].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[3].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

        count += 1
   
    # table.cell(0,0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

    # for row in table.rows:
    #     for cell in row.cells:
    #         # Center vertically
    #         cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #         # Center each paragraph horizontally
    #         for paragraph in cell.paragraphs:
    #             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table.columns[0].width = Inches(0.7)
    document.add_paragraph('')

    set_table_borders(table)


    # for cell in table.columns[1].cells:
    #     cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# Set table borders
def set_table_borders(table):
    tbl = table._tbl  # Get the underlying XML element
    tblBorders = OxmlElement('w:tblBorders')
    
    # Add all the different borders
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border size (in eighths of a point)
        border.set(qn('w:space'), '0')  # No spacing
        border.set(qn('w:color'), '000000')  # Black color
        tblBorders.append(border)
    
    # Apply the borders to the table
    tblPr = tbl.xpath('w:tblPr')[0]
    tblPr.append(tblBorders)

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

create_table(document, BNC1, COCA1, ALL_WORDS1)
create_table(document, BNC2, COCA2, ALL_WORDS2)
create_table(document, BNC3, COCA3, ALL_WORDS3)


#save the file
document.save('tableName.docx')

